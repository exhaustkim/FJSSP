"""FJSSP Research Workbench — Flask application entry point.

Start with:
    python app.py
Then open http://localhost:5000 in your browser.
"""

from __future__ import annotations

import inspect
import json
import math
import os
import threading
import time
import uuid
from copy import deepcopy
from pathlib import Path
from typing import Dict, List, Optional

import numpy as np
from flask import (
    Flask, Response, jsonify, render_template,
    request, send_file, stream_with_context,
)

# -- sim imports --
from sim.data.loader import (
    load_instance, load_instance_from_path,
    list_benchmarks, get_benchmark_stats,
    generate_due_dates, estimate_makespan,
    FAMILY_FOLDERS, CUSTOM_DIR,
)
from sim.core.simulator import FJSSPSimulator
from sim.rules.baseline import RULES as BASELINE_RULES
from sim.scenarios.s0_normal import S0Normal
from sim.scenarios.s1_part_delay import S1PartDelay
from sim.scenarios.s2_urgent_order import S2UrgentOrder
from sim.evaluation.metrics import summarise_runs, average_relative_improvement
from sim.evaluation.stats import wilcoxon_test
from sim.llm.evolution import EoHEvolution
from sim.llm.experience import ExperienceStore

# ---------------------------------------------------------------------------
# App setup
# ---------------------------------------------------------------------------

app = Flask(__name__, template_folder="ui/templates", static_folder="ui/static")
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB upload limit

RESULTS_DIR = Path("results")
RESULTS_DIR.mkdir(exist_ok=True)
CUSTOM_DIR.mkdir(parents=True, exist_ok=True)

experience_store = ExperienceStore(RESULTS_DIR / "experience_store.json")

# In-memory task registry  {task_id -> TaskState}
_tasks: Dict[str, dict] = {}
_tasks_lock = threading.Lock()

# Benchmark families exposed to the UI
ALLOWED_FAMILIES = {"brandimarte", "hurink", "behnke", "custom"}

# Global settings (overridable via /settings)
_settings: dict = {
    "ddt": 1.5,
    "n_seeds_final": 100,
    "n_seeds_eval": 5,
    "n_iter": 20,
    "pool_size": 15,
    "llm_model": "gpt-4o-mini",
    "llm_temperature": 0.8,
    "s1_affected_ratio": 0.20,
    "s1_delay_k": 1.0,
    "s2_due_date_factor": 0.5,
    "metric_weights": {"at": 1.0, "mit": 0.0, "ptj": 0.0},
    "time_unit_minutes": 1,   # 1 simulation unit = N minutes (for disruption time display)
}


# ---------------------------------------------------------------------------
# Task helpers
# ---------------------------------------------------------------------------

def _new_task(kind: str) -> str:
    tid = str(uuid.uuid4())
    with _tasks_lock:
        _tasks[tid] = {
            "id": tid, "kind": kind,
            "status": "queued",
            "progress": 0,
            "log": [],
            "result": None,
            "error": None,
        }
    return tid


def _task_log(tid: str, msg: str, progress: int = None):
    with _tasks_lock:
        t = _tasks.get(tid)
        if t:
            t["log"].append(msg)
            if progress is not None:
                t["progress"] = progress


def _task_done(tid: str, result):
    with _tasks_lock:
        t = _tasks.get(tid)
        if t:
            t["status"] = "done"
            t["progress"] = 100
            t["result"] = result


def _task_error(tid: str, err: str):
    with _tasks_lock:
        t = _tasks.get(tid)
        if t:
            t["status"] = "error"
            t["error"] = err


# ---------------------------------------------------------------------------
# Page routes
# ---------------------------------------------------------------------------

@app.route("/")
def page_home():
    return render_template("index.html", active="home")

@app.route("/benchmark-manager")
def page_benchmark():
    return render_template("benchmark_manager.html", active="benchmark")

@app.route("/scenario-manager")
def page_scenario():
    return render_template("scenario_manager.html", active="scenario", _settings=_settings)

@app.route("/rule-explorer")
def page_rule_explorer():
    return render_template("rule_explorer.html", active="rule_explorer")

@app.route("/evolution-center")
def page_evolution():
    return render_template("evolution_center.html", active="evolution")

@app.route("/simulation-results")
def page_results():
    return render_template("simulation_results.html", active="results")

@app.route("/stress-test")
def page_stress():
    return render_template("stress_test.html", active="stress")

@app.route("/similarity-analysis")
def page_similarity():
    return render_template("similarity_analysis.html", active="similarity")

@app.route("/memory-center")
def page_memory():
    return render_template("memory_center.html", active="memory")

@app.route("/human-vs-ai")
def page_human_vs_ai():
    return render_template("human_vs_ai.html", active="human_vs_ai")

@app.route("/report-generator")
def page_report():
    return render_template("report_generator.html", active="report")

@app.route("/event-translator")
def page_event_translator():
    return render_template("event_translator.html", active="event_translator")

@app.route("/settings")
def page_settings():
    return render_template("settings.html", active="settings")


# ---------------------------------------------------------------------------
# API — Benchmarks
# ---------------------------------------------------------------------------

@app.route("/api/benchmarks")
def api_list_benchmarks():
    all_bmarks = list_benchmarks()
    filtered = {k: v for k, v in all_bmarks.items() if k in ALLOWED_FAMILIES}
    return jsonify({"ok": True, "data": filtered})


@app.route("/api/benchmark/<family>/<name>")
def api_benchmark_stats(family, name):
    try:
        stats = get_benchmark_stats(name, family)
        return jsonify({"ok": True, "data": stats})
    except FileNotFoundError as e:
        return jsonify({"ok": False, "error": str(e)}), 404
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/benchmark/<family>/<name>/params")
def api_benchmark_params(family, name):
    """Return timing parameters needed for time-unit disruption controls."""
    try:
        import statistics as _stats
        jobs, n_m = load_instance(name, family)
        t_est = estimate_makespan(jobs, n_m)
        avg_min_pt = _stats.mean(j.get_total_min_pt() for j in jobs) if jobs else 1.0
        return jsonify({"ok": True, "data": {
            "t_est": round(t_est, 2),
            "avg_min_pt_per_job": round(avg_min_pt, 2),
            "n_jobs": len(jobs),
            "n_machines": n_m,
        }})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/benchmark/upload", methods=["POST"])
def api_upload_benchmark():
    if "file" not in request.files:
        return jsonify({"ok": False, "error": "No file"}), 400
    f = request.files["file"]
    if not f.filename.endswith(".json"):
        return jsonify({"ok": False, "error": "Only .json files accepted"}), 400
    save_path = CUSTOM_DIR / f.filename
    f.save(str(save_path))
    try:
        stats = get_benchmark_stats(save_path.stem, "custom")
        return jsonify({"ok": True, "data": stats})
    except Exception as e:
        save_path.unlink(missing_ok=True)
        return jsonify({"ok": False, "error": f"Invalid file: {e}"}), 400


# ---------------------------------------------------------------------------
# API — Tasks (progress streaming)
# ---------------------------------------------------------------------------

@app.route("/api/tasks/<tid>")
def api_task_status(tid):
    with _tasks_lock:
        t = _tasks.get(tid)
    if not t:
        return jsonify({"ok": False, "error": "unknown task"}), 404
    return jsonify({"ok": True, "data": t})


@app.route("/api/tasks/<tid>/stream")
def api_task_stream(tid):
    @stream_with_context
    def generate():
        while True:
            with _tasks_lock:
                t = dict(_tasks.get(tid, {}))
            if not t:
                yield f"data: {json.dumps({'error': 'not found'})}\n\n"
                break
            yield f"data: {json.dumps(t)}\n\n"
            if t.get("status") in ("done", "error"):
                break
            time.sleep(0.4)
    return Response(generate(), content_type="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ---------------------------------------------------------------------------
# API — Baselines
# ---------------------------------------------------------------------------

@app.route("/api/baselines/run", methods=["POST"])
def api_run_baselines():
    body     = request.json or {}
    name     = body.get("instance", "mk01")
    family   = body.get("family", "brandimarte")
    scenario = body.get("scenario", "S0")
    n_seeds  = int(body.get("n_seeds", _settings["n_seeds_final"]))
    ddt      = float(body.get("ddt", _settings["ddt"]))
    sp       = _extract_scenario_params(body)

    tid = _new_task("baselines")

    def worker():
        try:
            with _tasks_lock:
                _tasks[tid]["status"] = "running"
            jobs, n_m = load_instance(name, family)
            t_est     = estimate_makespan(jobs, n_m)
            scen_obj  = _build_scenario(scenario, **sp)

            def scen_factory(seed):
                return scen_obj.build_events(jobs, n_m, t_est, seed=seed)

            results = {}
            rule_ids = list(BASELINE_RULES.keys())
            for i, bid in enumerate(rule_ids):
                fn = BASELINE_RULES[bid]
                ats, ptjs, mits, makespans = _eval_rule_full(fn, jobs, n_m, scen_factory, n_seeds, 0, ddt)
                s = summarise_runs(ats)
                s["at_list"] = ats
                finite_ptjs = [p for p in ptjs if p < 200]
                finite_mits = [m for m in mits if m != float("inf")]
                finite_ms   = [m for m in makespans if m != float("inf")]
                s["ptj"]      = float(np.mean(finite_ptjs)) if finite_ptjs else 0.0
                s["mit"]      = float(np.mean(finite_mits)) if finite_mits else 0.0
                s["makespan"] = float(np.mean(finite_ms)) if finite_ms else 0.0
                results[bid] = s
                prog = int((i + 1) / len(rule_ids) * 100)
                _task_log(tid, f"{bid}: AT={s['mean']:.3f}±{s['std']:.3f}  PTJ={s['ptj']:.1f}%", prog)

            best_mean = min(v["mean"] for v in results.values())
            for bid in results:
                results[bid]["ari"] = average_relative_improvement(
                    best_mean, results[bid]["mean"])

            best_bid = min(results, key=lambda b: results[b]["mean"])
            best_ats = results[best_bid]["at_list"]
            for bid in results:
                if bid != best_bid:
                    stat, p = wilcoxon_test(best_ats, results[bid]["at_list"])
                    results[bid]["wilcoxon"] = {"stat": stat, "p": round(p, 4)}

            _save_results(results, f"baselines_{scenario}_{name}")
            _task_done(tid, {"results": results, "instance": name,
                             "scenario": scenario, "t_est": t_est})
        except Exception as exc:
            import traceback
            _task_error(tid, traceback.format_exc())

    threading.Thread(target=worker, daemon=True).start()
    return jsonify({"ok": True, "task_id": tid})


# ---------------------------------------------------------------------------
# API — Evolution (P1 / P2 / P3)
# ---------------------------------------------------------------------------

@app.route("/api/evolution/start", methods=["POST"])
def api_start_evolution():
    body     = request.json or {}
    name     = body.get("instance", "mk01")
    family   = body.get("family", "brandimarte")
    method   = body.get("method", "P2")
    scenario = body.get("scenario", "S0")
    n_seeds  = int(body.get("n_seeds", _settings["n_seeds_final"]))
    ddt      = float(body.get("ddt", _settings["ddt"]))
    n_iter   = int(body.get("n_iter", _settings["n_iter"]))
    sp       = _extract_scenario_params(body)

    use_ext = method in ("P2", "P3")
    use_exp = method == "P3"

    tid = _new_task("evolution")

    def worker():
        try:
            with _tasks_lock:
                _tasks[tid]["status"] = "running"

            jobs, n_m  = load_instance(name, family)
            t_est      = estimate_makespan(jobs, n_m)
            scen_obj   = _build_scenario(scenario, **sp)
            scen_tag   = _scenario_tag(scenario, **sp)
            scen_desc  = _scenario_desc(scenario, **sp)

            def scen_factory(seed):
                return scen_obj.build_events(jobs, n_m, t_est, seed=seed)

            def on_progress(iteration, best_at, msg):
                prog = int(iteration / n_iter * 85)
                _task_log(tid, msg, prog)
                with _tasks_lock:
                    _tasks[tid]["iter_best"] = (
                        _tasks[tid].get("iter_best", []) + [best_at]
                    )

            evo = EoHEvolution(
                jobs=jobs, n_machines=n_m,
                scenario_factory=scen_factory,
                scenario_tag=scen_tag,
                scenario_description=scen_desc,
                use_external_vars=use_ext,
                use_experience=use_exp,
                experience_store=experience_store,
                ddt=ddt,
                seed_offset=1000,
                progress_callback=on_progress,
            )
            evo.n_iter = n_iter
            result = evo.run()

            _task_log(tid, "Evaluating best rule on final seeds...", 90)
            best_ats = _eval_rule(
                result.best_rule.fn, jobs, n_m, scen_factory, n_seeds, 0, ddt
            )
            summary = summarise_runs(best_ats)
            summary["at_list"] = best_ats

            baseline_results = _load_results(f"baselines_{scenario}_{name}")
            if baseline_results:
                best_b_mean = min(v["mean"] for v in baseline_results.values())
                summary["ari"] = average_relative_improvement(best_b_mean, summary["mean"])

            out = {
                "method":             method,
                "instance":           name,
                "scenario":           scenario,
                "scenario_params":    sp,
                "best_rule":          result.best_rule.to_dict(),
                "iteration_best_at":  result.iteration_best_at,
                "genealogy":          result.genealogy,
                "generation_log":     result.generation_log,
                "summary":            summary,
                "pool_snapshots": [
                    [r.to_dict() for r in snap] for snap in result.pool_history
                ],
            }
            _save_results(out, f"evolution_{method}_{scenario}_{name}")
            _task_done(tid, out)
        except Exception:
            import traceback
            _task_error(tid, traceback.format_exc())

    threading.Thread(target=worker, daemon=True).start()
    return jsonify({"ok": True, "task_id": tid})


@app.route("/api/evolution/results")
def api_list_evolution_results():
    files = sorted(RESULTS_DIR.glob("evolution_*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    out = []
    for f in files[:20]:
        try:
            d = json.loads(f.read_text(encoding="utf-8"))
            out.append({
                "file": f.name,
                "method": d.get("method"),
                "instance": d.get("instance"),
                "scenario": d.get("scenario"),
                "best_at": d.get("summary", {}).get("mean"),
                "ari": d.get("summary", {}).get("ari"),
            })
        except Exception:
            pass
    return jsonify({"ok": True, "data": out})


@app.route("/api/evolution/result/<path:fname>")
def api_evolution_result(fname):
    path = RESULTS_DIR / fname
    if not path.exists():
        return jsonify({"ok": False, "error": "not found"}), 404
    return jsonify({"ok": True, "data": json.loads(path.read_text(encoding="utf-8"))})


# ---------------------------------------------------------------------------
# API — Simulation Results (stored)
# ---------------------------------------------------------------------------

@app.route("/api/results")
def api_list_results():
    files = sorted(RESULTS_DIR.glob("baselines_*.json"),
                   key=lambda p: p.stat().st_mtime, reverse=True)
    out = []
    for f in files[:30]:
        try:
            d = json.loads(f.read_text(encoding="utf-8"))
            out.append({"file": f.name, "data": d})
        except Exception:
            pass
    return jsonify({"ok": True, "data": out})


@app.route("/api/results/<path:fname>")
def api_get_result_file(fname):
    path = RESULTS_DIR / fname
    if not path.exists():
        return jsonify({"ok": False, "error": "not found"}), 404
    return jsonify({"ok": True, "data": json.loads(path.read_text(encoding="utf-8"))})


# ---------------------------------------------------------------------------
# API — Stress Test
# ---------------------------------------------------------------------------

@app.route("/api/stress-test/run", methods=["POST"])
def api_stress_test():
    body = request.json or {}
    name    = body.get("instance", "mk01")
    family  = body.get("family", "brandimarte")
    rules   = body.get("rules", list(BASELINE_RULES.keys()))
    n_seeds = int(body.get("n_seeds", 30))
    ddt     = float(body.get("ddt", _settings["ddt"]))

    tid = _new_task("stress_test")

    STRESS_SCENARIOS = [
        ("S0",  {}),
        ("S1",  {"s1_ratio": 0.10, "s1_k": 1.0}),
        ("S1",  {"s1_ratio": 0.20, "s1_k": 1.0}),
        ("S1",  {"s1_ratio": 0.40, "s1_k": 1.0}),
        ("S2",  {"s2_ddf": 0.3}),
        ("S2",  {"s2_ddf": 0.5}),
        ("S2",  {"s2_ddf": 1.0}),
    ]

    def worker():
        try:
            with _tasks_lock:
                _tasks[tid]["status"] = "running"

            jobs, n_m = load_instance(name, family)
            t_est = estimate_makespan(jobs, n_m)
            matrix = {}   # {rule_id: {scenario_label: mean_at}}
            total = len(rules) * len(STRESS_SCENARIOS)
            done = 0

            for bid in rules:
                if bid not in BASELINE_RULES:
                    continue
                fn = BASELINE_RULES[bid]
                matrix[bid] = {}
                for (stype, sparams) in STRESS_SCENARIOS:
                    s1r = sparams.get("s1_ratio", 0.20)
                    s1k = sparams.get("s1_k", 1.0)
                    s2d = sparams.get("s2_ddf", 0.5)
                    lbl = _scenario_label(stype, s1r, s1k, s2d)
                    scen_obj = _build_scenario(stype, s1r, s1k, s2d)

                    def sfactory(seed, _j=jobs, _n=n_m, _t=t_est, _s=scen_obj):
                        return _s.build_events(_j, _n, _t, seed=seed)

                    ats = _eval_rule(fn, jobs, n_m, sfactory, n_seeds, 0, ddt)
                    matrix[bid][lbl] = round(float(np.mean(ats)), 3)
                    done += 1
                    _task_log(tid, f"{bid}/{lbl}: AT={matrix[bid][lbl]:.3f}",
                              int(done / total * 100))

            scenario_labels = [_scenario_label(s, p.get("s1_ratio", 0.20),
                                               p.get("s1_k", 1.0), p.get("s2_ddf", 0.5))
                               for s, p in STRESS_SCENARIOS]
            out = {"matrix": matrix, "scenarios": scenario_labels, "instance": name}
            _save_results(out, f"stress_{name}")
            _task_done(tid, out)
        except Exception:
            import traceback
            _task_error(tid, traceback.format_exc())

    threading.Thread(target=worker, daemon=True).start()
    return jsonify({"ok": True, "task_id": tid})


@app.route("/api/stress-test/results")
def api_stress_results():
    files = sorted(RESULTS_DIR.glob("stress_*.json"),
                   key=lambda p: p.stat().st_mtime, reverse=True)
    out = []
    for f in files[:10]:
        try:
            out.append({"file": f.name,
                        "data": json.loads(f.read_text(encoding="utf-8"))})
        except Exception:
            pass
    return jsonify({"ok": True, "data": out})


# ---------------------------------------------------------------------------
# API — Similarity Analysis
# ---------------------------------------------------------------------------

@app.route("/api/similarity", methods=["POST"])
def api_similarity():
    body = request.json or {}
    include_llm_file = body.get("evolution_file")

    rules: Dict[str, str] = {}
    for bid, fn in BASELINE_RULES.items():
        try:
            rules[bid] = inspect.getsource(fn)
        except Exception:
            rules[bid] = f"# {bid}"

    if include_llm_file:
        path = RESULTS_DIR / include_llm_file
        if path.exists():
            d = json.loads(path.read_text(encoding="utf-8"))
            best = d.get("best_rule", {})
            if best.get("code"):
                rules[best.get("rule_id", "LLM_best")] = best["code"]
            for snap in d.get("pool_snapshots", [])[-1:]:
                for r in snap:
                    if r.get("code"):
                        rules[r["rule_id"]] = r["code"]

    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        from sklearn.decomposition import PCA
        from sklearn.manifold import TSNE
    except ImportError:
        return jsonify({"ok": False, "error": "scikit-learn not installed"}), 500

    ids = list(rules.keys())
    texts = [rules[i] for i in ids]
    vec = TfidfVectorizer(analyzer="char_wb", ngram_range=(3, 5))
    X = vec.fit_transform(texts).toarray()
    sim_matrix = cosine_similarity(X).tolist()

    pca2 = PCA(n_components=2)
    coords_pca = pca2.fit_transform(X).tolist()

    n = len(ids)
    perp = min(5, n - 1) if n > 2 else 1
    tsne_coords: List[List[float]] = []
    if n >= 3:
        tsne2 = TSNE(n_components=2, perplexity=perp, random_state=42)
        tsne_coords = tsne2.fit_transform(X).tolist()
    else:
        tsne_coords = coords_pca

    return jsonify({
        "ok": True,
        "data": {
            "ids": ids,
            "similarity_matrix": sim_matrix,
            "pca": coords_pca,
            "tsne": tsne_coords,
        }
    })


# ---------------------------------------------------------------------------
# API — Memory (P3 experience)
# ---------------------------------------------------------------------------

@app.route("/api/memory")
def api_list_memory():
    try:
        data = experience_store._load()
        return jsonify({"ok": True, "data": data})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/memory/<scenario_tag>")
def api_memory_by_tag(scenario_tag):
    try:
        entries = experience_store.retrieve_top(scenario_tag, k=50)
        return jsonify({"ok": True, "data": entries})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


# ---------------------------------------------------------------------------
# API — Rules (all available rules including LLM results)
# ---------------------------------------------------------------------------

@app.route("/api/rules")
def api_list_rules():
    rules = []
    for bid, fn in BASELINE_RULES.items():
        try:
            code = inspect.getsource(fn)
        except Exception:
            code = f"# {bid}"
        rules.append({
            "rule_id": bid,
            "code": code,
            "type": "baseline",
            "generation": 0,
            "parents": [],
            "operator": "BASELINE",
        })

    # Add best rules from evolution results
    for f in sorted(RESULTS_DIR.glob("evolution_*.json"),
                    key=lambda p: p.stat().st_mtime, reverse=True)[:5]:
        try:
            d = json.loads(f.read_text(encoding="utf-8"))
            best = d.get("best_rule", {})
            if best.get("code"):
                best["type"] = d.get("method", "LLM")
                best["file"] = f.name
                rules.append(best)
        except Exception:
            pass

    return jsonify({"ok": True, "data": rules})


# ---------------------------------------------------------------------------
# API — Report Generator
# ---------------------------------------------------------------------------

@app.route("/api/report/generate", methods=["POST"])
def api_generate_report():
    body = request.json or {}
    fmt  = body.get("format", "markdown")  # markdown | pdf | docx

    # Gather data
    sections = body.get("sections", ["benchmark", "scenario", "baselines", "evolution", "summary"])
    instance = body.get("instance", "")
    scenario = body.get("scenario", "S0")

    md_lines = ["# FJSSP Research Workbench — Experiment Report", ""]
    md_lines += [f"**Generated:** {time.strftime('%Y-%m-%d %H:%M:%S')}", ""]

    if "benchmark" in sections and instance:
        try:
            fam = body.get("family", "brandimarte")
            s = get_benchmark_stats(instance, fam)
            md_lines += [
                "## Benchmark",
                f"- Instance: **{s['name'].upper()}** ({s['family']})",
                f"- Jobs: {s['n_jobs']}  |  Machines: {s['n_machines']}",
                f"- Operations: {s['n_operations']}  |  Flexibility: {s['flexibility']:.2f}",
                "",
            ]
        except Exception:
            pass

    if "scenario" in sections:
        md_lines += [
            "## Scenario",
            f"- Type: **{scenario}**",
            f"- DDT: {body.get('ddt', _settings['ddt'])}",
            "",
        ]

    if "baselines" in sections:
        bdata = _load_results(f"baselines_{scenario}_{instance}")
        if bdata:
            md_lines += ["## 기본 규칙 성능 (Baseline Performance)", ""]
            md_lines += [
                "| 규칙 | AT 평균 | AT 표준편차 | 납기위반%(PTJ) | 기계유휴(MIT) | Makespan | ARI |",
                "|------|---------|-----------|--------------|------------|---------|-----|",
            ]
            for bid, v in sorted(bdata.items(), key=lambda x: x[1]["mean"]):
                md_lines.append(
                    f"| {bid} | {v.get('mean',0):.3f} | {v.get('std',0):.3f} "
                    f"| {v.get('ptj',0):.1f}% | {v.get('mit',0):.1f} "
                    f"| {v.get('makespan',0):.1f} | {v.get('ari',0):+.1f}% |"
                )
            md_lines.append("")
            best_bid = min(bdata, key=lambda b: bdata[b]["mean"])
            best_v = bdata[best_bid]
            md_lines += [
                f"> **최적 규칙**: {best_bid}  |  AT {best_v['mean']:.3f} ± {best_v['std']:.3f}"
                f"  |  납기위반 {best_v.get('ptj',0):.1f}%  |  Makespan {best_v.get('makespan',0):.1f}",
                "",
            ]

    if "evolution" in sections and instance:
        # Track the best evolution result for the summary impact section.
        best_evo_for_impact = None
        best_baseline_for_impact = None
        if "bdata" in locals() and bdata:
            try:
                best_bid_local = min(bdata, key=lambda b: bdata[b]["mean"])
                best_baseline_for_impact = bdata[best_bid_local]
            except Exception:
                pass

        for f in sorted(RESULTS_DIR.glob(f"evolution_*_{scenario}_{instance}.json"),
                        key=lambda p: p.stat().st_mtime, reverse=True)[:3]:
            try:
                d = json.loads(f.read_text(encoding="utf-8"))
                method = d.get("method", "?")
                s = d.get("summary", {})
                best = d.get("best_rule", {})
                md_lines += [
                    f"## {method} Evolution Result",
                    f"- Best AT: **{s.get('mean',0):.3f} ± {s.get('std',0):.3f}**",
                    f"- ARI vs best baseline: **{s.get('ari',0):+.1f}%**",
                    f"- Rule ID: `{best.get('rule_id','')}`",
                    "",
                    "**Generated Rule:**",
                    "```python",
                    (best.get("code") or "# (built-in rule)")[:500],
                    "```",
                    "",
                ]

                # NL explanation of the generated rule (LLM-powered).
                rule_code = best.get("code") or ""
                if rule_code and body.get("include_nl", True):
                    try:
                        from sim.llm.client import _client as _llm, MODEL as _MODEL
                        if _llm is not None:
                            nl_prompt = (
                                "다음 dispatching rule 코드를 한국어 자연어로 풀어 설명하세요. "
                                "JSON으로 답변하세요: {\"summary\":\"한 줄 요약\","
                                "\"intuition\":\"이 식이 왜 좋은지 직관 (2~3문장)\","
                                "\"when_strong\":\"강한 상황\",\"when_weak\":\"약한 상황\"}\n\n"
                                f"코드:\n```\n{rule_code[:1500]}\n```"
                            )
                            nl_resp = _llm.chat.completions.create(
                                model=_MODEL,
                                messages=[{"role": "user", "content": nl_prompt}],
                                temperature=0.2, max_tokens=500,
                                response_format={"type": "json_object"},
                            )
                            nl = json.loads(nl_resp.choices[0].message.content)
                            md_lines += [
                                "**🗣 규칙 자연어 설명 (이 규칙이 하는 일):**",
                                f"- **한 줄 요약:** {nl.get('summary','')}",
                                f"- **직관:** {nl.get('intuition','')}",
                                f"- **강한 상황:** {nl.get('when_strong','')}",
                                f"- **약한 상황:** {nl.get('when_weak','')}",
                                "",
                            ]
                    except Exception as _e:
                        md_lines += [f"_(자연어 설명 생성 실패: {_e})_", ""]

                # Track first (most recent) evolution result for the impact section.
                if best_evo_for_impact is None:
                    best_evo_for_impact = s
            except Exception:
                pass

        # ----- Concrete improvement-meaning section -----
        if best_evo_for_impact and best_baseline_for_impact:
            try:
                bb = best_baseline_for_impact
                ev = best_evo_for_impact
                # Reuse the same logic as /api/improvement/explain inline.
                b_at = float(bb.get("mean", 0));   p_at = float(ev.get("mean", 0))
                b_pt = float(bb.get("ptj", 0));    p_pt = float(ev.get("ptj", 0))
                b_mit = float(bb.get("mit", 0));   p_mit = float(ev.get("mit", 0))
                b_mk  = float(bb.get("makespan", 0)); p_mk = float(ev.get("makespan", 0))

                n_jobs = 0
                try:
                    n_jobs = int(get_benchmark_stats(instance, body.get("family", "brandimarte"))["n_jobs"])
                except Exception:
                    pass

                md_lines += ["## 🎯 개선의 구체적 의미", ""]
                if b_at > 1e-6:
                    at_pct = (b_at - p_at) / b_at * 100.0
                    at_delta = b_at - p_at
                    md_lines.append(
                        f"- **평균 납기 초과량(AT)**: {b_at:.2f} → {p_at:.2f} "
                        f"({at_pct:+.1f}%, 작업당 평균 약 {abs(at_delta):.2f}분 "
                        f"{'감소' if at_delta > 0 else '증가'})"
                    )
                if n_jobs and (b_pt or p_pt):
                    b_late = b_pt / 100.0 * n_jobs
                    p_late = p_pt / 100.0 * n_jobs
                    md_lines.append(
                        f"- **납기 위반 작업 수**: 약 {b_late:.1f}건 → {p_late:.1f}건 "
                        f"(전체 {n_jobs}건 중, {b_late - p_late:+.1f}건 더 납기 준수)"
                    )
                if b_mit and p_mit:
                    mit_pct = (b_mit - p_mit) / b_mit * 100.0
                    md_lines.append(
                        f"- **기계 유휴시간(MIT)**: {b_mit:.0f} → {p_mit:.0f} "
                        f"({mit_pct:+.1f}%) — 동일 자원 활용도"
                    )
                if b_mk and p_mk:
                    mk_pct = (b_mk - p_mk) / b_mk * 100.0
                    md_lines.append(
                        f"- **전체 완료시각(Makespan)**: {b_mk:.0f} → {p_mk:.0f} ({mk_pct:+.1f}%)"
                    )

                scen_note = {
                    "S1": "외부 충격(부품 지연)이 있는 상황에서도 납기 일정이 더 안정적으로 유지된다는 의미.",
                    "S2": "긴급 작업이 삽입된 상황에서 다른 작업들의 납기 손실 폭이 줄었다는 의미.",
                    "S0": "외부 충격이 없는 환경의 baseline 성능 — 기본 의사결정 자체가 더 좋아짐."
                }.get(scenario)
                if scen_note:
                    md_lines += ["", f"> 💡 **시나리오 해석**: {scen_note}"]
                md_lines.append("")
            except Exception:
                pass

    md_text = "\n".join(md_lines)

    if fmt == "markdown":
        tmp = RESULTS_DIR / "report.md"
        tmp.write_text(md_text, encoding="utf-8")
        return send_file(str(tmp), as_attachment=True, download_name="report.md",
                         mimetype="text/markdown")

    elif fmt == "pdf":
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=11)
            for line in md_text.replace("**", "").replace("`", "").splitlines():
                safe = line.encode("latin-1", "replace").decode("latin-1")
                pdf.cell(0, 7, safe, ln=True)
            tmp = RESULTS_DIR / "report.pdf"
            pdf.output(str(tmp))
            return send_file(str(tmp), as_attachment=True, download_name="report.pdf",
                             mimetype="application/pdf")
        except ImportError:
            return jsonify({"ok": False, "error": "fpdf2 not installed"}), 500

    elif fmt == "docx":
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("FJSSP Research Workbench Report", 0)
            for line in md_text.splitlines():
                if line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.strip():
                    doc.add_paragraph(line)
            tmp = RESULTS_DIR / "report.docx"
            doc.save(str(tmp))
            return send_file(str(tmp), as_attachment=True, download_name="report.docx",
                             mimetype="application/vnd.openxmlformats-officedocument"
                                      ".wordprocessingml.document")
        except ImportError:
            return jsonify({"ok": False, "error": "python-docx not installed"}), 500

    return jsonify({"ok": False, "error": "Unknown format"}), 400


# ---------------------------------------------------------------------------
# API — Action Recommendations
# ---------------------------------------------------------------------------

@app.route("/api/recommendations", methods=["POST"])
def api_recommendations():
    """Analyzes disruption settings and result data, returns Korean action recommendations."""
    body     = request.json or {}
    scenario = body.get("scenario", "S0")
    best_rule = body.get("best_rule", "")
    s1_ratio = float(body.get("s1_ratio", 0.20))
    s1_k     = float(body.get("s1_k", 1.0))
    s2_ddf   = float(body.get("s2_ddf", 0.5))
    results  = body.get("results", {})  # {rule_id: {mean, std, ptj, mit, makespan}}

    recs = []

    # ── S0: 정상 운영 ──
    if scenario == "S0":
        if best_rule in ("B3", "B10"):
            recs.append({"level": "success",
                         "text": f"✅ <strong>{best_rule}</strong>이(가) 정상 운영 최적 규칙입니다. "
                                 "현재 설정을 유지하며 운영하세요."})
        else:
            recs.append({"level": "info",
                         "text": f"현재 최적 규칙은 <strong>{best_rule}</strong>입니다. "
                                 "외부 교란이 없는 환경에서는 처리 시간 기반 규칙 "
                                 "<strong>B3 SPT</strong> · <strong>B10 ATCS</strong>가 "
                                 "처리량 극대화에 유리합니다."})
        recs.append({"level": "info",
                     "text": "S0 환경에서는 단순 규칙(B3, B6)도 충분합니다. "
                             "진화 센터 <strong>P1 방식</strong>으로 추가 개선 여지를 탐색해 보세요."})

    # ── S1: 부품 지연 ──
    elif scenario == "S1":
        if s1_ratio >= 0.4 or s1_k >= 2.0:
            recs.append({"level": "danger",
                         "text": f"⚠️ <strong>심각한 부품 지연</strong> — "
                                 f"{int(s1_ratio*100)}% 작업이 {s1_k:.1f}배 지연됩니다.<br>"
                                 "즉시 납기 여유(Slack)와 잔여 처리 시간을 동시에 고려하는 "
                                 "<strong>B6 PT+WINQ+SL</strong> 또는 "
                                 "<strong>B10 ATCS</strong>로 규칙을 전환하세요."})
            recs.append({"level": "warning",
                         "text": "부품 조달 일정을 재검토하고, 영향 작업을 사전 식별하여 "
                                 "우선 처리하는 전략을 권장합니다. "
                                 "가능하다면 대체 공급업체 확보 또는 안전 재고 확충을 검토하세요."})
        elif s1_ratio >= 0.2 or s1_k >= 1.0:
            recs.append({"level": "warning",
                         "text": f"🔶 <strong>중간 수준 부품 지연</strong> — "
                                 f"{int(s1_ratio*100)}% 작업 영향.<br>"
                                 "크리티컬 레이쇼 <strong>B4 CR</strong> 또는 "
                                 "<strong>B7 CR+SPT</strong> 규칙이 납기 준수에 유리합니다."})
            recs.append({"level": "info",
                         "text": "지연 발생 시 해당 작업의 납기 여유를 실시간 모니터링하고, "
                                 "가능한 빠른 기계에 재배정하는 것이 효과적입니다."})
        else:
            recs.append({"level": "info",
                         "text": f"ℹ️ <strong>경미한 부품 지연</strong> — "
                                 f"{int(s1_ratio*100)}% 영향, 현재 규칙으로 대응 가능합니다.<br>"
                                 "<strong>B6 PT+WINQ+SL</strong> 검토를 권장합니다."})
        recs.append({"level": "info",
                     "text": "💡 <strong>AI 진화 추천</strong>: S1 환경에서는 "
                             "<strong>P2 방식</strong> 진화를 통해 "
                             "<code>part_available_time</code> 변수를 활용하는 특화 규칙을 "
                             "자동 생성할 수 있습니다. 진화 센터에서 P2를 실행하세요."})

    # ── S2: 긴급 주문 ──
    elif scenario == "S2":
        if s2_ddf <= 0.3:
            recs.append({"level": "danger",
                         "text": f"🚨 <strong>매우 촉박한 긴급 주문</strong> — 납기 계수 {s2_ddf}.<br>"
                                 "<strong>B5 Urgency</strong> 규칙은 긴급 플래그를 1,000배 가중하여 "
                                 "즉시 우선 처리합니다. B5 사용을 강력 권장합니다."})
            recs.append({"level": "warning",
                         "text": "긴급 주문 발생 즉시 전담 기계를 확보하고, "
                                 "일반 작업을 일시 보류하는 선점(preemption) 전략을 검토하세요."})
        elif s2_ddf <= 0.5:
            recs.append({"level": "warning",
                         "text": f"🔶 <strong>촉박한 긴급 주문</strong> — 납기 계수 {s2_ddf}.<br>"
                                 "<strong>B5 Urgency</strong> 또는 P2 진화 규칙이 효과적입니다."})
            recs.append({"level": "info",
                         "text": "긴급 주문 처리 후 잔여 일반 작업의 납기 재조정(rescheduling)을 고려하세요."})
        else:
            recs.append({"level": "info",
                         "text": f"ℹ️ <strong>여유 있는 긴급 주문</strong> — 납기 계수 {s2_ddf}. "
                                 "대부분의 규칙이 처리 가능합니다. "
                                 "<strong>B10 ATCS</strong> 또는 <strong>B2 EDD</strong>를 권장합니다."})
        recs.append({"level": "info",
                     "text": "💡 <strong>AI 진화 추천</strong>: S2 환경에서는 "
                             "<strong>P2 방식</strong>이 <code>urgent_order_flag</code> 변수를 "
                             "학습하여 최적 규칙을 자동 생성합니다. 진화 센터에서 P2를 실행하세요."})

    # ── 성능 기반 권고 ──
    if results and best_rule:
        best_v = results.get(best_rule, {})
        ptj    = best_v.get("ptj", 0)
        makespan = best_v.get("makespan", 0)
        if ptj > 50:
            recs.append({"level": "danger",
                         "text": f"📊 <strong>높은 납기 위반율</strong>: 최적 규칙 {best_rule}에서도 "
                                 f"{ptj:.1f}% 지연 발생.<br>"
                                 "DDT 계수를 높이거나(납기 여유 증가), AI 진화(P2/P3)로 개선하세요."})
        elif ptj > 25:
            recs.append({"level": "warning",
                         "text": f"📊 납기 위반율 {ptj:.1f}% — 개선 여지가 있습니다. "
                                 "진화 센터에서 P2 규칙 생성을 시도하세요."})
        else:
            recs.append({"level": "success",
                         "text": f"📊 납기 위반율 {ptj:.1f}% — 양호한 수준입니다."})

    return jsonify({"ok": True, "data": recs})


# ---------------------------------------------------------------------------
# API — Settings
# ---------------------------------------------------------------------------

@app.route("/api/settings", methods=["GET"])
def api_get_settings():
    return jsonify({"ok": True, "data": _settings})


@app.route("/api/settings", methods=["POST"])
def api_save_settings():
    body = request.json or {}
    for key, val in body.items():
        if key in _settings:
            _settings[key] = val
    return jsonify({"ok": True, "data": _settings})


# ---------------------------------------------------------------------------
# API — Home stats
# ---------------------------------------------------------------------------

@app.route("/api/home/stats")
def api_home_stats():
    bmarks = list_benchmarks()
    n_instances = sum(len(v) for v in bmarks.values())
    n_baseline_rules = len(BASELINE_RULES)
    n_experiments = len(list(RESULTS_DIR.glob("baselines_*.json")))
    n_evolutions = len(list(RESULTS_DIR.glob("evolution_*.json")))

    recent = []
    for f in sorted(RESULTS_DIR.glob("*.json"),
                    key=lambda p: p.stat().st_mtime, reverse=True)[:5]:
        recent.append({"file": f.name,
                       "mtime": time.strftime("%Y-%m-%d %H:%M",
                                              time.localtime(f.stat().st_mtime))})
    return jsonify({
        "ok": True,
        "data": {
            "n_instances": n_instances,
            "n_baseline_rules": n_baseline_rules,
            "n_experiments": n_experiments,
            "n_evolutions": n_evolutions,
            "recent_files": recent,
        }
    })


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _build_scenario(stype,
                    s1_ratio=0.20, s1_k=1.0,
                    s1_timing_min=0.1, s1_timing_max=0.3,
                    s2_ddf=0.5, s2_arrival_min=0.2, s2_arrival_max=0.5,
                    s2_n_urgent=1):
    if stype == "S0":
        return S0Normal()
    if stype == "S1":
        return S1PartDelay(
            affected_ratio=s1_ratio, delay_k=s1_k,
            timing_min=s1_timing_min, timing_max=s1_timing_max,
        )
    if stype == "S2":
        return S2UrgentOrder(
            due_date_factor=s2_ddf,
            arrival_min=s2_arrival_min, arrival_max=s2_arrival_max,
            n_urgent_jobs=s2_n_urgent,
        )
    raise ValueError(f"Unknown scenario: {stype}")


def _extract_scenario_params(body: dict) -> dict:
    """Request body에서 시나리오 관련 파라미터를 일괄 추출."""
    s = _settings
    return dict(
        s1_ratio       = float(body.get("s1_ratio",       s["s1_affected_ratio"])),
        s1_k           = float(body.get("s1_k",           s["s1_delay_k"])),
        s1_timing_min  = float(body.get("s1_timing_min",  0.1)),
        s1_timing_max  = float(body.get("s1_timing_max",  0.3)),
        s2_ddf         = float(body.get("s2_ddf",         s["s2_due_date_factor"])),
        s2_arrival_min = float(body.get("s2_arrival_min", 0.2)),
        s2_arrival_max = float(body.get("s2_arrival_max", 0.5)),
        s2_n_urgent    = int(body.get("s2_n_urgent",      1)),
    )


def _scenario_tag(stype, s1_ratio=0.20, s1_k=1.0, s2_ddf=0.5, **kw) -> str:
    if stype == "S1":
        return f"S1_r{s1_ratio}_k{s1_k}"
    if stype == "S2":
        return f"S2_d{s2_ddf}"
    return "S0"


def _scenario_desc(stype, s1_ratio=0.20, s1_k=1.0,
                   s1_timing_min=0.1, s1_timing_max=0.3,
                   s2_ddf=0.5, s2_arrival_min=0.2, s2_arrival_max=0.5,
                   s2_n_urgent=1, **kw) -> str:
    if stype == "S0":
        return "정상 운영 (외부 교란 없음)."
    if stype == "S1":
        return (f"부품 지연: {s1_ratio*100:.0f}% 작업 영향, "
                f"지연 배수 k={s1_k}, "
                f"트리거 시점 [{s1_timing_min}~{s1_timing_max}]×T_est.")
    if stype == "S2":
        return (f"긴급 주문 {s2_n_urgent}건, "
                f"납기 계수={s2_ddf}, "
                f"도착 시점 [{s2_arrival_min}~{s2_arrival_max}]×T_est.")
    return ""


def _scenario_label(stype, s1_ratio=0.20, s1_k=1.0, s2_ddf=0.5, **kw) -> str:
    if stype == "S0":
        return "S0"
    if stype == "S1":
        return f"S1 {int(s1_ratio*100)}% k={s1_k}"
    if stype == "S2":
        return f"S2 ddf={s2_ddf}"
    return stype


def _eval_rule(fn, jobs, n_machines, scen_factory, n_seeds, seed_offset, ddt):
    """Evaluate a rule, return list of AT values only."""
    ats = []
    for seed in range(seed_offset, seed_offset + n_seeds):
        j = generate_due_dates(deepcopy(jobs), ddt=ddt, seed=seed)
        events = scen_factory(seed)
        sim = FJSSPSimulator(j, n_machines, fn, scenario_events=events)
        try:
            ats.append(sim.run().at)
        except Exception:
            ats.append(float("inf"))
    return ats


def _eval_rule_full(fn, jobs, n_machines, scen_factory, n_seeds, seed_offset, ddt):
    """Evaluate a rule, return (ats, ptjs, mits, makespans) across seeds."""
    ats, ptjs, mits, makespans = [], [], [], []
    for seed in range(seed_offset, seed_offset + n_seeds):
        j = generate_due_dates(deepcopy(jobs), ddt=ddt, seed=seed)
        events = scen_factory(seed)
        sim = FJSSPSimulator(j, n_machines, fn, scenario_events=events)
        try:
            r = sim.run()
            ats.append(r.at)
            ptjs.append(r.ptj)
            mits.append(r.mit_total)
            makespans.append(r.makespan)
        except Exception:
            ats.append(float("inf"))
            ptjs.append(100.0)
            mits.append(float("inf"))
            makespans.append(float("inf"))
    return ats, ptjs, mits, makespans


def _save_results(data: dict, tag: str):
    path = RESULTS_DIR / f"{tag}.json"
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(data, fh, indent=2, ensure_ascii=False, default=str)


def _load_results(tag: str) -> Optional[dict]:
    path = RESULTS_DIR / f"{tag}.json"
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return None


# ---------------------------------------------------------------------------
# API — LLM Narrative Report  (Human vs AI 서술형 분석)
# ---------------------------------------------------------------------------

@app.route("/api/report/narrative", methods=["POST"])
def api_report_narrative():
    """Generate LLM narrative analysis comparing human baselines vs AI-evolved rule."""
    from sim.llm.client import _client as _llm, MODEL as _MODEL

    body = request.json or {}
    baseline_data = body.get("baseline_data", {})
    evo_data      = body.get("evo_data", {})
    instance      = body.get("instance", "")
    scenario      = body.get("scenario", "S0")

    if not baseline_data:
        return jsonify({"ok": False, "error": "기본 규칙 데이터가 없습니다."}), 400
    if _llm is None:
        return jsonify({"ok": False, "error": "OpenAI API 키가 설정되지 않았습니다."}), 503

    sorted_b = sorted(baseline_data.items(), key=lambda x: x[1].get("mean", 9999))
    best_b, worst_b = sorted_b[0], sorted_b[-1]
    spread_pct = ((worst_b[1].get("mean", 0) - best_b[1].get("mean", 0))
                  / max(worst_b[1].get("mean", 0.001), 0.001) * 100)

    baseline_rows = "\n".join(
        f"  {bid}: AT={v.get('mean',0):.3f}±{v.get('std',0):.3f},"
        f" 납기위반={v.get('ptj',0):.1f}%, ARI={v.get('ari',0):+.1f}%"
        for bid, v in sorted_b
    )

    evo_block = "(AI 진화 결과 없음 — 기본 규칙만 분석)"
    if evo_data:
        best = evo_data.get("best_rule", {})
        s    = evo_data.get("summary", {})
        evo_block = (
            f"방식: {evo_data.get('method','LLM')}  |  규칙: {best.get('rule_id','—')}\n"
            f"  AT={s.get('mean',0):.3f}±{s.get('std',0):.3f},"
            f" 납기위반={s.get('ptj',0):.1f}%,"
            f" ARI={s.get('ari',0):+.1f}%,"
            f" 발견={best.get('generation','—')}세대,"
            f" 생성방식={best.get('operator_kr', best.get('operator','—'))}"
        )

    scen_desc = {"S0": "정상 운영", "S1": "부품 지연 교란", "S2": "긴급 주문 교란"}.get(scenario, scenario)

    prompt = f"""당신은 제조 스케줄링 AI 연구 보고서 작성 전문가입니다.
FJSSP(유연 작업장 스케줄링 문제) 실험 결과를 분석하여 한국어 서술형 보고서를 작성하세요.

== 실험 개요 ==
벤치마크: {instance.upper() if instance else '미지정'}  |  시나리오: {scenario} ({scen_desc})

== 기본 규칙(인간 설계) 성능 — AT 오름차순 ==
{baseline_rows}
최적: {best_b[0]} (AT {best_b[1].get('mean',0):.3f})  |  최악: {worst_b[0]} (AT {worst_b[1].get('mean',0):.3f})  |  개선폭: {spread_pct:.1f}%

== AI 진화 결과 ==
{evo_block}

== 작성 지침 ==
아래 4개 항목을 각각 제목 포함 단락으로 작성하세요. 수치는 반드시 인용하고, 전문적이되 실무자도 이해할 수 있는 문체로 작성하세요.

**1. 기본 규칙 성능 분석**
B1~B10 규칙 간 성능 분포, 최적/최악 규칙의 알고리즘 특성, 시나리오 적합성을 2~3문장으로 서술.

**2. {'인간 vs AI 비교 분석' if evo_data else '규칙 전략 심화 분석'}**
{'AI 진화 규칙이 기본 규칙 대비 ARI 수치를 활용해 얼마나 개선했는지, 핵심 차별점을 3~4문장으로 비교.' if evo_data else '최적·차선 규칙의 알고리즘 차이와 시나리오별 우위 요인을 3~4문장으로 분석.'}

**3. {'AI 진화 규칙 실무 적용 평가' if evo_data else '시나리오 대응 전략'}**
{'AI 규칙의 코드 복잡도, 일반화 가능성, 현장 도입 시 고려사항을 2~3문장으로 평가.' if evo_data else '해당 시나리오 특성에 맞는 규칙 운용 전략을 2~3문장으로 제안.'}

**4. 결론 및 권장 사항**
핵심 발견사항과 실무 적용 권장 사항을 2~3문장으로 요약."""

    try:
        resp = _llm.chat.completions.create(
            model=_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=1500,
        )
        return jsonify({"ok": True, "narrative": resp.choices[0].message.content})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


# ---------------------------------------------------------------------------
# API — Natural-language explanation of a dispatching rule
# ---------------------------------------------------------------------------

@app.route("/api/rule/explain", methods=["POST"])
def api_rule_explain():
    """규칙(코드)을 한국어 자연어 설명으로 변환.

    Input:
        code         : str — rule body (Python or DSL expression)
        rule_id      : str (optional) — for context, e.g. 'B3' or 'P3-G7-M1'
        context      : dict (optional) — scenario, metrics, etc.

    Output:
        {ok: True, summary: 한 줄 요약,
         details: {variables_used, when_strong, when_weak, intuition}}
    """
    from sim.llm.client import _client as _llm, MODEL as _MODEL

    body     = request.json or {}
    code     = (body.get("code") or "").strip()
    rule_id  = body.get("rule_id", "")
    ctx      = body.get("context", {}) or {}

    if not code:
        return jsonify({"ok": False, "error": "code 필드가 비어 있습니다."}), 400
    if _llm is None:
        return jsonify({"ok": False, "error": "OpenAI API 키가 설정되지 않았습니다."}), 503

    scenario = ctx.get("scenario", "")
    at_mean  = ctx.get("at_mean")
    ari      = ctx.get("ari")
    ctx_line = []
    if scenario:
        ctx_line.append(f"시나리오: {scenario}")
    if at_mean is not None:
        ctx_line.append(f"AT: {at_mean:.3f}")
    if ari is not None:
        ctx_line.append(f"ARI: {ari:+.1f}%")
    ctx_block = " | ".join(ctx_line) or "(컨텍스트 없음)"

    prompt = f"""당신은 제조 스케줄링 휴리스틱을 비전공자도 이해할 수 있게 설명하는 전문가입니다.
아래 dispatching rule 코드를 한국어 자연어로 풀어주세요. 수학식이 아닌 평이한 문장으로.

== 컨텍스트 ==
규칙 ID: {rule_id or '미지정'}
{ctx_block}

== 규칙 코드 ==
```
{code[:2000]}
```

== 작성 지침 ==
JSON 형식으로 답하되, 코드 블록이나 설명 없이 순수 JSON만 출력하세요.
{{
  "summary": "이 규칙이 하는 일을 한 문장으로 (e.g., '납기가 임박하고 처리시간이 짧은 작업을 우선 처리한다')",
  "variables_used": ["release_time", "due_date", "processing_time" 같이 핵심 변수 3~5개],
  "intuition": "왜 이런 식이 좋은지 직관적인 설명 (2~3문장, 실무자 톤)",
  "when_strong": "이 규칙이 강한 상황 (e.g., '납기 압박이 큰 환경, 작업 길이가 비슷할 때')",
  "when_weak": "이 규칙이 약할 수 있는 상황 (e.g., '긴급 주문이 자주 들어오면 우선순위가 흐트러짐')"
}}"""

    try:
        resp = _llm.chat.completions.create(
            model=_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=600,
            response_format={"type": "json_object"},
        )
        import json as _json
        parsed = _json.loads(resp.choices[0].message.content)
        return jsonify({"ok": True, **parsed})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


# ---------------------------------------------------------------------------
# API — Concrete "impact" interpretation of an ARI improvement
# ---------------------------------------------------------------------------

@app.route("/api/improvement/explain", methods=["POST"])
def api_improvement_explain():
    """ARI/PTJ/MIT 개선치를 실무 관점의 구체적 의미로 변환.

    Input:
        baseline: {mean, ptj, mit, makespan}    # AT_mean, PTJ%, MIT, makespan
        proposed: {mean, ptj, mit, makespan}
        n_jobs (optional): 인스턴스 작업 수
        scenario (optional): S0/S1/S2

    Output:
        {ok, impact_lines: [...], headline}
    """
    body = request.json or {}
    b = body.get("baseline", {}) or {}
    p = body.get("proposed", {}) or {}
    n_jobs   = body.get("n_jobs", 0) or 0
    scenario = body.get("scenario", "")

    def _f(d, k, default=0.0):
        try: return float(d.get(k, default))
        except Exception: return default

    b_at = _f(b, "mean");  p_at = _f(p, "mean")
    b_pt = _f(b, "ptj");   p_pt = _f(p, "ptj")
    b_mit = _f(b, "mit");  p_mit = _f(p, "mit")
    b_mk  = _f(b, "makespan");  p_mk  = _f(p, "makespan")

    impact = []
    headline = ""

    if b_at > 1e-6:
        at_pct = (b_at - p_at) / b_at * 100.0
        at_delta = b_at - p_at
        sign = "감소" if at_delta > 0 else "증가"
        impact.append(
            f"평균 납기 초과량(AT): **{b_at:.2f} → {p_at:.2f}** "
            f"({at_pct:+.1f}%, 작업당 평균 약 {abs(at_delta):.2f}분 {sign})"
        )
        headline = f"AT가 baseline 대비 {at_pct:+.1f}% {'개선' if at_delta > 0 else '악화'}"

    if n_jobs and (b_pt or p_pt):
        b_late = b_pt / 100.0 * n_jobs
        p_late = p_pt / 100.0 * n_jobs
        impact.append(
            f"납기 위반 작업 수: 약 **{b_late:.1f}건 → {p_late:.1f}건** "
            f"(전체 {n_jobs}건 중, {b_late - p_late:+.1f}건 더 납기 준수)"
        )

    if b_mit and p_mit:
        mit_pct = (b_mit - p_mit) / b_mit * 100.0 if b_mit else 0
        impact.append(
            f"기계 유휴시간(MIT): **{b_mit:.0f} → {p_mit:.0f}** "
            f"({mit_pct:+.1f}%) — 동일 자원으로 다른 의사결정"
        )

    if b_mk and p_mk:
        mk_pct = (b_mk - p_mk) / b_mk * 100.0
        impact.append(
            f"전체 완료시각(Makespan): **{b_mk:.0f} → {p_mk:.0f}** ({mk_pct:+.1f}%)"
        )

    interp = []
    if headline:
        interp.append(f"### 한 줄 요약\n{headline}.")

    if impact:
        interp.append("### 구체적 영향")
        interp.extend(f"- {line}" for line in impact)

    # Scenario-specific reading
    if scenario == "S1":
        interp.append(
            "### 시나리오 의미 (S1: 부품 지연)\n"
            "AT 감소는 외부 충격이 발생한 상황에서도 납기 일정을 더 안정적으로 지킨다는 뜻. "
            "PTJ가 같이 줄었다면 다수 작업의 큰 지연이 아닌 분산 효과가 작동."
        )
    elif scenario == "S2":
        interp.append(
            "### 시나리오 의미 (S2: 긴급 주문)\n"
            "긴급 작업이 삽입된 상황에서 다른 작업들의 납기 손실 폭이 줄어듦. "
            "MIT가 비슷하면 머신 가동률은 그대로이면서 더 영리하게 스케줄링한 것."
        )
    elif scenario == "S0":
        interp.append(
            "### 시나리오 의미 (S0: 정상 운영)\n"
            "외부 충격이 없는 상황의 baseline 성능. 이 환경에서 개선되면 기본 dispatching "
            "의사결정 자체가 더 좋아졌다는 의미."
        )

    return jsonify({
        "ok": True,
        "headline": headline,
        "impact_lines": impact,
        "interpretation_md": "\n\n".join(interp),
    })


# ---------------------------------------------------------------------------
# API — Supply Chain Event → Simulation Parameter Translation
# ---------------------------------------------------------------------------

@app.route("/api/event/translate", methods=["POST"])
def api_event_translate():
    """Translate a real-world supply chain disruption event into FJSSP simulation parameters."""
    import re as _re
    from sim.llm.client import _client as _llm, MODEL as _MODEL

    body          = request.json or {}
    event_types   = body.get("event_types", [])
    delay_days    = float(body.get("delay_days", 0))
    affected_parts = body.get("affected_parts", "")
    severity      = body.get("severity", "medium")
    additional    = body.get("additional_info", "")

    if _llm is None:
        return jsonify({"ok": False, "error": "OpenAI API 키가 설정되지 않았습니다."}), 503

    sev_kr = {"low": "경미", "medium": "중간", "high": "심각"}.get(severity, "중간")

    prompt = f"""당신은 공급망 이벤트를 FJSSP 생산 스케줄링 시뮬레이션 파라미터로 변환하는 전문가입니다.

[FJSSP 시뮬레이션 시나리오]
S1 (부품 지연): 일부 작업의 부품 도착이 지연됨
  - s1_ratio: 영향받는 작업 비율 (0.05~1.0)
  - s1_k: 지연량 배수 — 지연시간 = avg_min_pt × s1_k (0.3~5.0)
  - s1_timing_min/max: 지연 이벤트 발생 타이밍 (전체 T_est 대비 비율, 0.0~1.0)

S2 (긴급 주문): 스케줄 도중 긴급 작업이 삽입됨
  - s2_ddf: 납기 여유 배수 (0.1=매우 촉박 ~ 2.0=여유)
  - s2_n_urgent: 긴급 주문 건수 (1~5)
  - s2_arrival_min/max: 주문 도착 타이밍 (0.0~1.0)

[입력 이벤트]
유형: {', '.join(event_types) if event_types else '미지정'}
예상 지연: {delay_days}일
영향 부품: {affected_parts or '미지정'}
심각도: {sev_kr}
추가 정보: {additional or '없음'}

[변환 규칙]
- 항만 폐쇄/공급업체 화재/파업/지정학적 리스크 → S1
- 품질 문제 → S1 (재작업 지연)
- 긴급 주문 → S2
- 복합 이벤트 (S1 계열 + 긴급 주문) → S1+S2
- 지연 일수가 많고 심각도가 높을수록 s1_k 증가, s1_ratio 증가
- 긴급 주문 납기가 촉박할수록 s2_ddf 감소

반드시 아래 JSON 형식만 반환하세요:
{{
  "primary_scenario": "S1" 또는 "S2" 또는 "S1+S2",
  "s1_params": {{"s1_ratio": 숫자, "s1_k": 숫자, "s1_timing_min": 숫자, "s1_timing_max": 숫자}},
  "s2_params": {{"s2_ddf": 숫자, "s2_n_urgent": 정수, "s2_arrival_min": 숫자, "s2_arrival_max": 숫자}},
  "rationale": {{
    "scenario_reason": "시나리오 선택 이유 (1~2문장)",
    "s1_ratio_reason": "비율 설정 근거 (1문장, S1 없으면 null)",
    "s1_k_reason": "배수 설정 근거 (1문장, S1 없으면 null)",
    "s2_reason": "S2 파라미터 설정 근거 (1문장, S2 없으면 null)"
  }},
  "impact_summary": "이벤트가 생산 스케줄에 미치는 영향 요약 (2~3문장)"
}}"""

    try:
        resp = _llm.chat.completions.create(
            model=_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=900,
            response_format={"type": "json_object"},
        )
        raw = resp.choices[0].message.content.strip()
        result = json.loads(raw)
        return jsonify({"ok": True, "data": result})
    except json.JSONDecodeError:
        # fallback: extract JSON block
        try:
            m = _re.search(r'\{.*\}', raw, _re.DOTALL)
            if m:
                result = json.loads(m.group())
                return jsonify({"ok": True, "data": result})
        except Exception:
            pass
        return jsonify({"ok": False, "error": "LLM 응답 JSON 파싱 오류", "raw": raw[:500]}), 500
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys
    port = int(sys.argv[1]) if len(sys.argv) > 1 else 5000
    print(f"\n FJSSP Research Workbench")
    print(f" http://localhost:{port}\n")
    app.run(debug=True, port=port, threaded=True)
